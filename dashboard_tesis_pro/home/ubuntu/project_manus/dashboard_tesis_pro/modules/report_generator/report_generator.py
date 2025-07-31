#!/usr/bin/env python3
"""
Dashboard Tesis Pro - Módulo Generador de Informes Profesionales

Este módulo proporciona funcionalidades completas para:
- Generación automática de reportes de análisis
- Plantillas profesionales personalizables
- Exportación en múltiples formatos (PDF, HTML, Word, Markdown)
- Integración de gráficos, tablas y estadísticas
- Reportes ejecutivos y técnicos
"""

import os
import json
import base64
from datetime import datetime
from typing import Dict, List, Any, Optional
import pandas as pd
import numpy as np
from jinja2 import Environment, FileSystemLoader, Template
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
import io
import zipfile

# Importaciones para diferentes formatos
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ProfessionalReportGenerator:
    """Generador de informes profesionales para análisis de datos"""
    
    def __init__(self, templates_dir: str = None):
        self.templates_dir = templates_dir or os.path.join(os.path.dirname(__file__), 'templates')
        self.assets_dir = os.path.join(os.path.dirname(__file__), 'assets')
        self.exports_dir = os.path.join(os.path.dirname(__file__), 'exports')
        
        # Crear directorios si no existen
        os.makedirs(self.templates_dir, exist_ok=True)
        os.makedirs(self.assets_dir, exist_ok=True)
        os.makedirs(self.exports_dir, exist_ok=True)
        
        # Configurar Jinja2
        self.jinja_env = Environment(
            loader=FileSystemLoader(self.templates_dir),
            autoescape=True
        )
        
        # Registrar filtros personalizados
        self.jinja_env.filters['format_number'] = self._format_number
        self.jinja_env.filters['format_percentage'] = self._format_percentage
        self.jinja_env.filters['format_date'] = self._format_date
        
        # Configuración por defecto
        self.default_config = {
            'company_name': 'Dashboard Tesis Pro',
            'company_logo': None,
            'author': 'Sistema de Análisis Automático',
            'language': 'es',
            'theme': 'professional',
            'color_scheme': {
                'primary': '#2196F3',
                'secondary': '#1976D2',
                'accent': '#FF9800',
                'success': '#4CAF50',
                'warning': '#FF9800',
                'error': '#F44336'
            }
        }
    
    def generate_comprehensive_report(
        self,
        data: pd.DataFrame,
        analysis_results: Dict[str, Any],
        report_config: Dict[str, Any] = None,
        template_name: str = 'comprehensive_analysis',
        output_format: str = 'html'
    ) -> Dict[str, Any]:
        """
        Generar reporte completo de análisis
        
        Args:
            data: DataFrame con los datos analizados
            analysis_results: Resultados de análisis estadísticos
            report_config: Configuración del reporte
            template_name: Nombre de la plantilla a usar
            output_format: Formato de salida ('html', 'pdf', 'docx', 'markdown')
        
        Returns:
            Dict con información del reporte generado
        """
        
        # Combinar configuración
        config = {**self.default_config, **(report_config or {})}
        
        # Preparar contexto para la plantilla
        context = self._prepare_report_context(data, analysis_results, config)
        
        # Generar reporte según el formato
        if output_format == 'html':
            return self._generate_html_report(context, template_name, config)
        elif output_format == 'pdf':
            return self._generate_pdf_report(context, template_name, config)
        elif output_format == 'docx':
            return self._generate_docx_report(context, template_name, config)
        elif output_format == 'markdown':
            return self._generate_markdown_report(context, template_name, config)
        else:
            raise ValueError(f"Formato no soportado: {output_format}")
    
    def generate_executive_summary(
        self,
        data: pd.DataFrame,
        key_findings: List[str],
        recommendations: List[str],
        config: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """Generar resumen ejecutivo"""
        
        context = {
            'report_type': 'executive_summary',
            'title': 'Resumen Ejecutivo - Análisis de Datos',
            'generated_at': datetime.now(),
            'data_summary': {
                'total_records': len(data),
                'total_columns': len(data.columns),
                'data_quality_score': self._calculate_data_quality_score(data),
                'analysis_period': self._get_analysis_period(data)
            },
            'key_findings': key_findings,
            'recommendations': recommendations,
            'config': {**self.default_config, **(config or {})}
        }
        
        return self._generate_html_report(context, 'executive_summary', context['config'])
    
    def generate_technical_report(
        self,
        data: pd.DataFrame,
        statistical_tests: Dict[str, Any],
        visualizations: List[Dict[str, Any]],
        methodology: str,
        config: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """Generar reporte técnico detallado"""
        
        context = {
            'report_type': 'technical_report',
            'title': 'Reporte Técnico - Análisis Estadístico Detallado',
            'generated_at': datetime.now(),
            'methodology': methodology,
            'data_description': self._generate_data_description(data),
            'statistical_tests': statistical_tests,
            'visualizations': visualizations,
            'technical_details': self._generate_technical_details(data),
            'config': {**self.default_config, **(config or {})}
        }
        
        return self._generate_html_report(context, 'technical_report', context['config'])
    
    def generate_dashboard_report(
        self,
        data: pd.DataFrame,
        charts: List[Dict[str, Any]],
        metrics: Dict[str, Any],
        config: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """Generar reporte tipo dashboard interactivo"""
        
        context = {
            'report_type': 'dashboard_report',
            'title': 'Dashboard de Análisis - Vista Interactiva',
            'generated_at': datetime.now(),
            'key_metrics': metrics,
            'charts': charts,
            'data_overview': self._generate_data_overview(data),
            'config': {**self.default_config, **(config or {})}
        }
        
        return self._generate_html_report(context, 'dashboard_report', context['config'])
    
    def _prepare_report_context(
        self,
        data: pd.DataFrame,
        analysis_results: Dict[str, Any],
        config: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Preparar contexto completo para la plantilla"""
        
        context = {
            'title': analysis_results.get('title', 'Reporte de Análisis de Datos'),
            'generated_at': datetime.now(),
            'config': config,
            
            # Información de datos
            'data_info': {
                'filename': analysis_results.get('filename', 'datos_analizados'),
                'total_records': len(data),
                'total_columns': len(data.columns),
                'memory_usage_mb': data.memory_usage(deep=True).sum() / 1024**2,
                'data_types': data.dtypes.value_counts().to_dict(),
                'missing_values_total': data.isnull().sum().sum(),
                'missing_percentage': (data.isnull().sum().sum() / (len(data) * len(data.columns))) * 100
            },
            
            # Estadísticas descriptivas
            'descriptive_stats': self._generate_descriptive_statistics(data),
            
            # Resultados de análisis
            'analysis_results': analysis_results,
            
            # Visualizaciones
            'visualizations': self._prepare_visualizations(data, analysis_results),
            
            # Conclusiones y recomendaciones
            'conclusions': self._generate_conclusions(data, analysis_results),
            'recommendations': self._generate_recommendations(data, analysis_results),
            
            # Metadatos
            'metadata': {
                'report_version': '2.0.0',
                'generator': 'Dashboard Tesis Pro',
                'analysis_engine': 'Python/Pandas/Plotly',
                'export_timestamp': datetime.now().isoformat()
            }
        }
        
        return context
    
    def _generate_html_report(
        self,
        context: Dict[str, Any],
        template_name: str,
        config: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generar reporte en formato HTML"""
        
        try:
            # Cargar plantilla
            template = self._get_or_create_template(template_name, 'html')
            
            # Renderizar HTML
            html_content = template.render(**context)
            
            # Generar nombre de archivo
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"reporte_{template_name}_{timestamp}.html"
            filepath = os.path.join(self.exports_dir, filename)
            
            # Guardar archivo
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return {
                'success': True,
                'format': 'html',
                'filename': filename,
                'filepath': filepath,
                'size_bytes': len(html_content.encode('utf-8')),
                'content': html_content
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'format': 'html'
            }
    
    def _generate_pdf_report(
        self,
        context: Dict[str, Any],
        template_name: str,
        config: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generar reporte en formato PDF"""
        
        if not WEASYPRINT_AVAILABLE:
            return {
                'success': False,
                'error': 'WeasyPrint no está disponible para generación de PDF',
                'format': 'pdf'
            }
        
        try:
            # Primero generar HTML
            html_result = self._generate_html_report(context, template_name, config)
            
            if not html_result['success']:
                return html_result
            
            # Convertir HTML a PDF
            html_content = html_result['content']
            
            # Configurar CSS para PDF
            css_content = self._get_pdf_css(config)
            
            # Generar PDF
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"reporte_{template_name}_{timestamp}.pdf"
            filepath = os.path.join(self.exports_dir, filename)
            
            HTML(string=html_content).write_pdf(
                filepath,
                stylesheets=[CSS(string=css_content)]
            )
            
            # Obtener tamaño del archivo
            file_size = os.path.getsize(filepath)
            
            return {
                'success': True,
                'format': 'pdf',
                'filename': filename,
                'filepath': filepath,
                'size_bytes': file_size
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'format': 'pdf'
            }
    
    def _generate_docx_report(
        self,
        context: Dict[str, Any],
        template_name: str,
        config: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generar reporte en formato Word (DOCX)"""
        
        if not DOCX_AVAILABLE:
            return {
                'success': False,
                'error': 'python-docx no está disponible para generación de Word',
                'format': 'docx'
            }
        
        try:
            # Crear documento
            doc = Document()
            
            # Configurar estilos
            self._setup_docx_styles(doc, config)
            
            # Agregar contenido
            self._add_docx_content(doc, context)
            
            # Generar nombre de archivo
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"reporte_{template_name}_{timestamp}.docx"
            filepath = os.path.join(self.exports_dir, filename)
            
            # Guardar documento
            doc.save(filepath)
            
            # Obtener tamaño del archivo
            file_size = os.path.getsize(filepath)
            
            return {
                'success': True,
                'format': 'docx',
                'filename': filename,
                'filepath': filepath,
                'size_bytes': file_size
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'format': 'docx'
            }
    
    def _generate_markdown_report(
        self,
        context: Dict[str, Any],
        template_name: str,
        config: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generar reporte en formato Markdown"""
        
        try:
            # Cargar plantilla Markdown
            template = self._get_or_create_template(template_name, 'md')
            
            # Renderizar Markdown
            markdown_content = template.render(**context)
            
            # Generar nombre de archivo
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"reporte_{template_name}_{timestamp}.md"
            filepath = os.path.join(self.exports_dir, filename)
            
            # Guardar archivo
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            return {
                'success': True,
                'format': 'markdown',
                'filename': filename,
                'filepath': filepath,
                'size_bytes': len(markdown_content.encode('utf-8')),
                'content': markdown_content
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'format': 'markdown'
            }
    
    def _get_or_create_template(self, template_name: str, format_ext: str) -> Template:
        """Obtener plantilla existente o crear una por defecto"""
        
        template_file = f"{template_name}.{format_ext}"
        template_path = os.path.join(self.templates_dir, template_file)
        
        # Si la plantilla existe, cargarla
        if os.path.exists(template_path):
            return self.jinja_env.get_template(template_file)
        
        # Si no existe, crear plantilla por defecto
        default_template = self._create_default_template(template_name, format_ext)
        
        # Guardar plantilla por defecto
        with open(template_path, 'w', encoding='utf-8') as f:
            f.write(default_template)
        
        # Recargar entorno Jinja2 y devolver plantilla
        self.jinja_env = Environment(
            loader=FileSystemLoader(self.templates_dir),
            autoescape=True
        )
        
        return self.jinja_env.get_template(template_file)
    
    def _create_default_template(self, template_name: str, format_ext: str) -> str:
        """Crear plantilla por defecto según el tipo"""
        
        if format_ext == 'html':
            return self._create_default_html_template(template_name)
        elif format_ext == 'md':
            return self._create_default_markdown_template(template_name)
        else:
            raise ValueError(f"Formato de plantilla no soportado: {format_ext}")
    
    def _create_default_html_template(self, template_name: str) -> str:
        """Crear plantilla HTML por defecto"""
        
        if template_name == 'comprehensive_analysis':
            return """
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ title }}</title>
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            background-color: #f8f9fa;
            color: #333;
        }
        .container {
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
        }
        .header {
            text-align: center;
            border-bottom: 3px solid {{ config.color_scheme.primary }};
            padding-bottom: 20px;
            margin-bottom: 30px;
        }
        .header h1 {
            color: {{ config.color_scheme.primary }};
            margin: 0;
            font-size: 2.5em;
        }
        .header .subtitle {
            color: #666;
            font-size: 1.2em;
            margin-top: 10px;
        }
        .section {
            margin: 30px 0;
            padding: 20px;
            border-left: 4px solid {{ config.color_scheme.primary }};
            background: #f8f9fa;
        }
        .section h2 {
            color: {{ config.color_scheme.primary }};
            margin-top: 0;
        }
        .metric-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin: 20px 0;
        }
        .metric-card {
            background: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            text-align: center;
        }
        .metric-value {
            font-size: 2em;
            font-weight: bold;
            color: {{ config.color_scheme.primary }};
        }
        .metric-label {
            color: #666;
            margin-top: 5px;
        }
        .table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }
        .table th, .table td {
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }
        .table th {
            background-color: {{ config.color_scheme.primary }};
            color: white;
        }
        .table tr:nth-child(even) {
            background-color: #f2f2f2;
        }
        .visualization {
            text-align: center;
            margin: 30px 0;
        }
        .footer {
            text-align: center;
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            color: #666;
            font-size: 0.9em;
        }
        .highlight {
            background-color: #fff3cd;
            border: 1px solid #ffeaa7;
            padding: 15px;
            border-radius: 5px;
            margin: 15px 0;
        }
        .success {
            background-color: #d4edda;
            border: 1px solid #c3e6cb;
            color: #155724;
        }
        .warning {
            background-color: #fff3cd;
            border: 1px solid #ffeaa7;
            color: #856404;
        }
        .error {
            background-color: #f8d7da;
            border: 1px solid #f5c6cb;
            color: #721c24;
        }
    </style>
</head>
<body>
    <div class="container">
        <!-- Header -->
        <div class="header">
            <h1>{{ title }}</h1>
            <div class="subtitle">
                Generado el {{ generated_at.strftime('%d de %B de %Y a las %H:%M') }}
            </div>
            <div class="subtitle">
                {{ config.company_name }} - {{ config.author }}
            </div>
        </div>

        <!-- Resumen de Datos -->
        <div class="section">
            <h2>📊 Resumen de Datos</h2>
            <div class="metric-grid">
                <div class="metric-card">
                    <div class="metric-value">{{ data_info.total_records | format_number }}</div>
                    <div class="metric-label">Registros</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{{ data_info.total_columns }}</div>
                    <div class="metric-label">Columnas</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{{ data_info.memory_usage_mb | round(2) }} MB</div>
                    <div class="metric-label">Memoria</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{{ data_info.missing_percentage | format_percentage }}</div>
                    <div class="metric-label">Datos Faltantes</div>
                </div>
            </div>
            
            <p><strong>Archivo analizado:</strong> {{ data_info.filename }}</p>
            <p><strong>Tipos de datos:</strong></p>
            <ul>
                {% for dtype, count in data_info.data_types.items() %}
                <li>{{ dtype }}: {{ count }} columnas</li>
                {% endfor %}
            </ul>
        </div>

        <!-- Estadísticas Descriptivas -->
        {% if descriptive_stats %}
        <div class="section">
            <h2>📈 Estadísticas Descriptivas</h2>
            
            {% if descriptive_stats.numeric %}
            <h3>Variables Numéricas</h3>
            <table class="table">
                <thead>
                    <tr>
                        <th>Variable</th>
                        <th>Media</th>
                        <th>Mediana</th>
                        <th>Desv. Estándar</th>
                        <th>Mínimo</th>
                        <th>Máximo</th>
                    </tr>
                </thead>
                <tbody>
                    {% for var, stats in descriptive_stats.numeric.items() %}
                    <tr>
                        <td>{{ var }}</td>
                        <td>{{ stats.mean | round(4) }}</td>
                        <td>{{ stats.median | round(4) }}</td>
                        <td>{{ stats.std | round(4) }}</td>
                        <td>{{ stats.min | round(4) }}</td>
                        <td>{{ stats.max | round(4) }}</td>
                    </tr>
                    {% endfor %}
                </tbody>
            </table>
            {% endif %}

            {% if descriptive_stats.categorical %}
            <h3>Variables Categóricas</h3>
            <table class="table">
                <thead>
                    <tr>
                        <th>Variable</th>
                        <th>Valores Únicos</th>
                        <th>Más Frecuente</th>
                        <th>Frecuencia</th>
                    </tr>
                </thead>
                <tbody>
                    {% for var, stats in descriptive_stats.categorical.items() %}
                    <tr>
                        <td>{{ var }}</td>
                        <td>{{ stats.unique_count }}</td>
                        <td>{{ stats.most_frequent }}</td>
                        <td>{{ stats.frequency }}</td>
                    </tr>
                    {% endfor %}
                </tbody>
            </table>
            {% endif %}
        </div>
        {% endif %}

        <!-- Resultados de Análisis -->
        {% if analysis_results %}
        <div class="section">
            <h2>🔬 Resultados de Análisis</h2>
            
            {% for analysis_type, results in analysis_results.items() %}
            {% if analysis_type not in ['title', 'filename'] %}
            <h3>{{ analysis_type | title }}</h3>
            
            {% if results is mapping %}
                {% for key, value in results.items() %}
                <p><strong>{{ key }}:</strong> {{ value }}</p>
                {% endfor %}
            {% else %}
                <p>{{ results }}</p>
            {% endif %}
            {% endif %}
            {% endfor %}
        </div>
        {% endif %}

        <!-- Visualizaciones -->
        {% if visualizations %}
        <div class="section">
            <h2>📊 Visualizaciones</h2>
            {% for viz in visualizations %}
            <div class="visualization">
                <h3>{{ viz.title }}</h3>
                {% if viz.image_base64 %}
                <img src="data:image/png;base64,{{ viz.image_base64 }}" alt="{{ viz.title }}" style="max-width: 100%; height: auto;">
                {% endif %}
                {% if viz.description %}
                <p>{{ viz.description }}</p>
                {% endif %}
            </div>
            {% endfor %}
        </div>
        {% endif %}

        <!-- Conclusiones -->
        {% if conclusions %}
        <div class="section">
            <h2>💡 Conclusiones</h2>
            {% for conclusion in conclusions %}
            <div class="highlight success">
                {{ conclusion }}
            </div>
            {% endfor %}
        </div>
        {% endif %}

        <!-- Recomendaciones -->
        {% if recommendations %}
        <div class="section">
            <h2>🎯 Recomendaciones</h2>
            {% for recommendation in recommendations %}
            <div class="highlight warning">
                {{ recommendation }}
            </div>
            {% endfor %}
        </div>
        {% endif %}

        <!-- Footer -->
        <div class="footer">
            <p>Reporte generado por {{ metadata.generator }} v{{ metadata.report_version }}</p>
            <p>Motor de análisis: {{ metadata.analysis_engine }}</p>
            <p>Timestamp: {{ metadata.export_timestamp }}</p>
        </div>
    </div>
</body>
</html>
            """
        
        # Plantillas adicionales para otros tipos de reporte
        elif template_name == 'executive_summary':
            return """
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ title }}</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }
        .header { text-align: center; border-bottom: 2px solid #2196F3; padding-bottom: 20px; }
        .section { margin: 30px 0; }
        .metric { background: #f8f9fa; padding: 15px; margin: 10px 0; border-left: 4px solid #2196F3; }
        .finding { background: #e8f5e8; padding: 15px; margin: 10px 0; border-radius: 5px; }
        .recommendation { background: #fff3e0; padding: 15px; margin: 10px 0; border-radius: 5px; }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ title }}</h1>
        <p>{{ generated_at.strftime('%d de %B de %Y') }}</p>
    </div>
    
    <div class="section">
        <h2>📊 Resumen de Datos</h2>
        <div class="metric">
            <strong>Registros analizados:</strong> {{ data_summary.total_records | format_number }}
        </div>
        <div class="metric">
            <strong>Variables:</strong> {{ data_summary.total_columns }}
        </div>
        <div class="metric">
            <strong>Calidad de datos:</strong> {{ data_summary.data_quality_score }}/100
        </div>
    </div>
    
    <div class="section">
        <h2>🔍 Hallazgos Clave</h2>
        {% for finding in key_findings %}
        <div class="finding">{{ finding }}</div>
        {% endfor %}
    </div>
    
    <div class="section">
        <h2>🎯 Recomendaciones</h2>
        {% for recommendation in recommendations %}
        <div class="recommendation">{{ recommendation }}</div>
        {% endfor %}
    </div>
</body>
</html>
            """
        
        else:
            # Plantilla genérica
            return """
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ title }}</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }
        .header { text-align: center; border-bottom: 2px solid #2196F3; padding-bottom: 20px; }
        .section { margin: 30px 0; }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ title }}</h1>
        <p>{{ generated_at.strftime('%d de %B de %Y a las %H:%M') }}</p>
    </div>
    
    <div class="section">
        <h2>Contenido del Reporte</h2>
        <p>Este es un reporte generado automáticamente.</p>
    </div>
</body>
</html>
            """
    
    def _create_default_markdown_template(self, template_name: str) -> str:
        """Crear plantilla Markdown por defecto"""
        
        return """
# {{ title }}

**Generado:** {{ generated_at.strftime('%d de %B de %Y a las %H:%M') }}  
**Autor:** {{ config.author }}  
**Organización:** {{ config.company_name }}

---

## 📊 Resumen de Datos

- **Archivo:** {{ data_info.filename }}
- **Registros:** {{ data_info.total_records | format_number }}
- **Columnas:** {{ data_info.total_columns }}
- **Memoria:** {{ data_info.memory_usage_mb | round(2) }} MB
- **Datos faltantes:** {{ data_info.missing_percentage | format_percentage }}

### Tipos de Datos

{% for dtype, count in data_info.data_types.items() %}
- **{{ dtype }}:** {{ count }} columnas
{% endfor %}

---

{% if descriptive_stats %}
## 📈 Estadísticas Descriptivas

{% if descriptive_stats.numeric %}
### Variables Numéricas

| Variable | Media | Mediana | Desv. Estándar | Mínimo | Máximo |
|----------|-------|---------|----------------|--------|--------|
{% for var, stats in descriptive_stats.numeric.items() %}
| {{ var }} | {{ stats.mean | round(4) }} | {{ stats.median | round(4) }} | {{ stats.std | round(4) }} | {{ stats.min | round(4) }} | {{ stats.max | round(4) }} |
{% endfor %}
{% endif %}

{% if descriptive_stats.categorical %}
### Variables Categóricas

| Variable | Valores Únicos | Más Frecuente | Frecuencia |
|----------|----------------|---------------|------------|
{% for var, stats in descriptive_stats.categorical.items() %}
| {{ var }} | {{ stats.unique_count }} | {{ stats.most_frequent }} | {{ stats.frequency }} |
{% endfor %}
{% endif %}
{% endif %}

---

{% if conclusions %}
## 💡 Conclusiones

{% for conclusion in conclusions %}
> {{ conclusion }}

{% endfor %}
{% endif %}

{% if recommendations %}
## 🎯 Recomendaciones

{% for recommendation in recommendations %}
- {{ recommendation }}
{% endfor %}
{% endif %}

---

**Reporte generado por {{ metadata.generator }} v{{ metadata.report_version }}**  
**Motor de análisis:** {{ metadata.analysis_engine }}  
**Timestamp:** {{ metadata.export_timestamp }}
        """
    
    # Métodos auxiliares para preparar datos
    def _generate_descriptive_statistics(self, data: pd.DataFrame) -> Dict[str, Any]:
        """Generar estadísticas descriptivas organizadas"""
        
        stats = {
            'numeric': {},
            'categorical': {}
        }
        
        # Estadísticas numéricas
        numeric_cols = data.select_dtypes(include=[np.number]).columns
        for col in numeric_cols:
            col_data = data[col].dropna()
            if len(col_data) > 0:
                stats['numeric'][col] = {
                    'mean': col_data.mean(),
                    'median': col_data.median(),
                    'std': col_data.std(),
                    'min': col_data.min(),
                    'max': col_data.max(),
                    'count': len(col_data),
                    'missing': data[col].isnull().sum()
                }
        
        # Estadísticas categóricas
        categorical_cols = data.select_dtypes(include=['object', 'category']).columns
        for col in categorical_cols:
            col_data = data[col].dropna()
            if len(col_data) > 0:
                value_counts = col_data.value_counts()
                stats['categorical'][col] = {
                    'unique_count': col_data.nunique(),
                    'most_frequent': value_counts.index[0] if len(value_counts) > 0 else 'N/A',
                    'frequency': value_counts.iloc[0] if len(value_counts) > 0 else 0,
                    'count': len(col_data),
                    'missing': data[col].isnull().sum()
                }
        
        return stats
    
    def _prepare_visualizations(self, data: pd.DataFrame, analysis_results: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Preparar visualizaciones para el reporte"""
        
        visualizations = []
        
        # Generar gráficos básicos automáticamente
        numeric_cols = data.select_dtypes(include=[np.number]).columns.tolist()
        categorical_cols = data.select_dtypes(include=['object', 'category']).columns.tolist()
        
        # Histograma de la primera variable numérica
        if len(numeric_cols) > 0:
            col = numeric_cols[0]
            fig = px.histogram(data, x=col, title=f"Distribución de {col}")
            fig.update_layout(template="plotly_white")
            
            # Convertir a imagen base64
            img_bytes = fig.to_image(format="png", width=800, height=500)
            img_base64 = base64.b64encode(img_bytes).decode()
            
            visualizations.append({
                'title': f"Distribución de {col}",
                'type': 'histogram',
                'image_base64': img_base64,
                'description': f"Histograma mostrando la distribución de valores en la variable {col}"
            })
        
        # Gráfico de barras de la primera variable categórica
        if len(categorical_cols) > 0:
            col = categorical_cols[0]
            value_counts = data[col].value_counts().head(10)  # Top 10
            
            fig = px.bar(
                x=value_counts.index,
                y=value_counts.values,
                title=f"Frecuencia de {col}"
            )
            fig.update_layout(template="plotly_white")
            
            # Convertir a imagen base64
            img_bytes = fig.to_image(format="png", width=800, height=500)
            img_base64 = base64.b64encode(img_bytes).decode()
            
            visualizations.append({
                'title': f"Frecuencia de {col}",
                'type': 'bar',
                'image_base64': img_base64,
                'description': f"Gráfico de barras mostrando la frecuencia de valores en {col}"
            })
        
        # Matriz de correlación si hay suficientes variables numéricas
        if len(numeric_cols) >= 2:
            corr_matrix = data[numeric_cols].corr()
            
            fig = px.imshow(
                corr_matrix,
                title="Matriz de Correlación",
                color_continuous_scale="RdBu_r",
                aspect="auto"
            )
            fig.update_layout(template="plotly_white")
            
            # Convertir a imagen base64
            img_bytes = fig.to_image(format="png", width=800, height=600)
            img_base64 = base64.b64encode(img_bytes).decode()
            
            visualizations.append({
                'title': "Matriz de Correlación",
                'type': 'heatmap',
                'image_base64': img_base64,
                'description': "Matriz de correlación entre variables numéricas"
            })
        
        return visualizations
    
    def _generate_conclusions(self, data: pd.DataFrame, analysis_results: Dict[str, Any]) -> List[str]:
        """Generar conclusiones automáticas basadas en los datos"""
        
        conclusions = []
        
        # Análisis de tamaño de dataset
        if len(data) > 10000:
            conclusions.append(f"El dataset contiene {len(data):,} registros, lo que constituye una muestra robusta para análisis estadísticos.")
        elif len(data) > 1000:
            conclusions.append(f"El dataset contiene {len(data):,} registros, proporcionando una base sólida para el análisis.")
        else:
            conclusions.append(f"El dataset contiene {len(data):,} registros. Se recomienda cautela en la generalización de resultados.")
        
        # Análisis de calidad de datos
        missing_pct = (data.isnull().sum().sum() / (len(data) * len(data.columns))) * 100
        if missing_pct < 5:
            conclusions.append("La calidad de los datos es excelente, con menos del 5% de valores faltantes.")
        elif missing_pct < 15:
            conclusions.append(f"La calidad de los datos es buena, con {missing_pct:.1f}% de valores faltantes.")
        else:
            conclusions.append(f"Se detectaron {missing_pct:.1f}% de valores faltantes. Se recomienda implementar estrategias de imputación.")
        
        # Análisis de variables
        numeric_cols = len(data.select_dtypes(include=[np.number]).columns)
        categorical_cols = len(data.select_dtypes(include=['object', 'category']).columns)
        
        if numeric_cols > categorical_cols:
            conclusions.append("El dataset es predominantemente numérico, ideal para análisis estadísticos y modelado predictivo.")
        elif categorical_cols > numeric_cols:
            conclusions.append("El dataset contiene principalmente variables categóricas, apropiado para análisis de frecuencias y asociaciones.")
        else:
            conclusions.append("El dataset presenta un balance entre variables numéricas y categóricas.")
        
        return conclusions
    
    def _generate_recommendations(self, data: pd.DataFrame, analysis_results: Dict[str, Any]) -> List[str]:
        """Generar recomendaciones automáticas"""
        
        recommendations = []
        
        # Recomendaciones basadas en valores faltantes
        missing_cols = data.columns[data.isnull().any()].tolist()
        if missing_cols:
            recommendations.append(f"Implementar estrategias de manejo de valores faltantes en las columnas: {', '.join(missing_cols[:3])}{'...' if len(missing_cols) > 3 else ''}")
        
        # Recomendaciones basadas en outliers
        numeric_cols = data.select_dtypes(include=[np.number]).columns
        outlier_cols = []
        for col in numeric_cols:
            Q1 = data[col].quantile(0.25)
            Q3 = data[col].quantile(0.75)
            IQR = Q3 - Q1
            outliers = data[(data[col] < Q1 - 1.5 * IQR) | (data[col] > Q3 + 1.5 * IQR)][col]
            if len(outliers) > len(data) * 0.05:  # Más del 5% son outliers
                outlier_cols.append(col)
        
        if outlier_cols:
            recommendations.append(f"Revisar y tratar valores atípicos en: {', '.join(outlier_cols[:3])}")
        
        # Recomendaciones de análisis adicionales
        if len(numeric_cols) >= 2:
            recommendations.append("Realizar análisis de correlación y regresión para identificar relaciones entre variables.")
        
        if len(data.select_dtypes(include=['object', 'category']).columns) > 0:
            recommendations.append("Considerar análisis de asociación entre variables categóricas usando pruebas chi-cuadrado.")
        
        if len(data) > 1000:
            recommendations.append("El tamaño del dataset permite aplicar técnicas de machine learning para análisis predictivo.")
        
        return recommendations
    
    # Métodos auxiliares para formateo
    def _format_number(self, value):
        """Formatear números con separadores de miles"""
        if pd.isna(value):
            return 'N/A'
        return f"{value:,.0f}"
    
    def _format_percentage(self, value):
        """Formatear porcentajes"""
        if pd.isna(value):
            return 'N/A'
        return f"{value:.1f}%"
    
    def _format_date(self, value):
        """Formatear fechas"""
        if isinstance(value, datetime):
            return value.strftime('%d/%m/%Y')
        return str(value)
    
    def _calculate_data_quality_score(self, data: pd.DataFrame) -> int:
        """Calcular puntuación de calidad de datos"""
        score = 100
        
        # Penalizar por valores faltantes
        missing_pct = (data.isnull().sum().sum() / (len(data) * len(data.columns))) * 100
        score -= missing_pct
        
        # Penalizar por duplicados
        duplicate_pct = (data.duplicated().sum() / len(data)) * 100
        score -= duplicate_pct * 2
        
        # Bonificar por diversidad de tipos de datos
        unique_types = len(data.dtypes.unique())
        if unique_types > 2:
            score += 5
        
        return max(0, min(100, int(score)))
    
    def _get_analysis_period(self, data: pd.DataFrame) -> str:
        """Obtener período de análisis si hay columnas de fecha"""
        date_cols = data.select_dtypes(include=['datetime64']).columns
        
        if len(date_cols) > 0:
            date_col = date_cols[0]
            min_date = data[date_col].min()
            max_date = data[date_col].max()
            return f"{min_date.strftime('%d/%m/%Y')} - {max_date.strftime('%d/%m/%Y')}"
        
        return "No especificado"
    
    def _generate_data_description(self, data: pd.DataFrame) -> Dict[str, Any]:
        """Generar descripción detallada de los datos"""
        return {
            'shape': data.shape,
            'dtypes': data.dtypes.value_counts().to_dict(),
            'memory_usage': data.memory_usage(deep=True).sum(),
            'missing_values': data.isnull().sum().to_dict(),
            'duplicate_rows': data.duplicated().sum()
        }
    
    def _generate_technical_details(self, data: pd.DataFrame) -> Dict[str, Any]:
        """Generar detalles técnicos para reporte técnico"""
        return {
            'data_types_detail': data.dtypes.to_dict(),
            'memory_usage_by_column': data.memory_usage(deep=True).to_dict(),
            'null_counts': data.isnull().sum().to_dict(),
            'unique_counts': data.nunique().to_dict(),
            'sample_values': {col: data[col].dropna().head(3).tolist() for col in data.columns}
        }
    
    def _generate_data_overview(self, data: pd.DataFrame) -> Dict[str, Any]:
        """Generar vista general para dashboard"""
        numeric_cols = data.select_dtypes(include=[np.number]).columns
        categorical_cols = data.select_dtypes(include=['object', 'category']).columns
        
        return {
            'total_records': len(data),
            'total_columns': len(data.columns),
            'numeric_columns': len(numeric_cols),
            'categorical_columns': len(categorical_cols),
            'missing_values_total': data.isnull().sum().sum(),
            'duplicate_rows': data.duplicated().sum(),
            'memory_usage_mb': data.memory_usage(deep=True).sum() / 1024**2
        }
    
    def _get_pdf_css(self, config: Dict[str, Any]) -> str:
        """Obtener CSS optimizado para PDF"""
        return f"""
        @page {{
            size: A4;
            margin: 2cm;
        }}
        
        body {{
            font-family: 'DejaVu Sans', Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.4;
            color: #333;
        }}
        
        .container {{
            max-width: none;
            margin: 0;
            padding: 0;
            box-shadow: none;
        }}
        
        .header {{
            border-bottom: 2pt solid {config['color_scheme']['primary']};
            margin-bottom: 20pt;
        }}
        
        .section {{
            page-break-inside: avoid;
            margin: 15pt 0;
        }}
        
        .metric-grid {{
            display: block;
        }}
        
        .metric-card {{
            display: inline-block;
            width: 23%;
            margin: 1%;
            vertical-align: top;
        }}
        
        .table {{
            font-size: 9pt;
        }}
        
        .visualization img {{
            max-width: 100%;
            height: auto;
        }}
        """
    
    def _setup_docx_styles(self, doc: Document, config: Dict[str, Any]):
        """Configurar estilos para documento Word"""
        
        # Estilo para títulos
        title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Estilo para subtítulos
        subtitle_style = doc.styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_style.font.name = 'Calibri'
        subtitle_style.font.size = Pt(14)
        subtitle_style.font.bold = True
        
        # Estilo para texto normal
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
    
    def _add_docx_content(self, doc: Document, context: Dict[str, Any]):
        """Agregar contenido al documento Word"""
        
        # Título
        title = doc.add_paragraph(context['title'], style='CustomTitle')
        
        # Fecha
        date_para = doc.add_paragraph(f"Generado el {context['generated_at'].strftime('%d de %B de %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Resumen de datos
        doc.add_paragraph('Resumen de Datos', style='CustomSubtitle')
        
        data_info = context['data_info']
        doc.add_paragraph(f"Archivo: {data_info['filename']}")
        doc.add_paragraph(f"Registros: {data_info['total_records']:,}")
        doc.add_paragraph(f"Columnas: {data_info['total_columns']}")
        doc.add_paragraph(f"Memoria: {data_info['memory_usage_mb']:.2f} MB")
        
        # Agregar más contenido según sea necesario
        if context.get('conclusions'):
            doc.add_paragraph('Conclusiones', style='CustomSubtitle')
            for conclusion in context['conclusions']:
                doc.add_paragraph(f"• {conclusion}")
        
        if context.get('recommendations'):
            doc.add_paragraph('Recomendaciones', style='CustomSubtitle')
            for recommendation in context['recommendations']:
                doc.add_paragraph(f"• {recommendation}")

# Función de utilidad para crear instancia del generador
def create_report_generator(templates_dir: str = None) -> ProfessionalReportGenerator:
    """Crear instancia del generador de reportes"""
    return ProfessionalReportGenerator(templates_dir)

